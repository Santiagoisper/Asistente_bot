"""Convert Markdown to DOCX (shareholder reports). Usage: python scripts/md-to-docx.py <input.md> [output.docx]"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    run_elem = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    r_pr.append(color)
    r_pr.append(underline)
    run_elem.append(r_pr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    run_elem.append(text_elem)
    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)


def parse_inline(text: str, paragraph) -> None:
    pattern = re.compile(
        r'(\*\*[^*]+\*\*|\[[^\]]+\]\([^)]+\)|`[^`]+`|\*[^*]+\*)'
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        chunk = m.group(0)
        if chunk.startswith('**') and chunk.endswith('**'):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        elif chunk.startswith('`') and chunk.endswith('`'):
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        elif chunk.startswith('['):
            lm = re.match(r'\[([^\]]+)\]\(([^)]+)\)', chunk)
            if lm:
                add_hyperlink(paragraph, lm.group(1), lm.group(2))
        elif chunk.startswith('*') and chunk.endswith('*'):
            run = paragraph.add_run(chunk[1:-1])
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith('|') and s.endswith('|') and '|' in s[1:-1]


def is_separator_row(line: str) -> bool:
    return bool(re.match(r'^\|[\s\-:|]+\|\s*$', line.strip()))


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and is_table_row(lines[i]):
        if not is_separator_row(lines[i]):
            cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
            rows.append(cells)
        i += 1
    return rows, i


class MdToDocx:
    def __init__(self, doc: Document) -> None:
        self.doc = doc
        self._setup_styles()

    def _setup_styles(self) -> None:
        normal = self.doc.styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(11)
        for level, size in [(1, 22), (2, 16), (3, 13), (4, 12)]:
            style = self.doc.styles[f'Heading {level}']
            style.font.name = 'Calibri'
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    def convert(self, md: str) -> None:
        lines = md.splitlines()
        i = 0
        in_code = False
        code_lang = ''
        code_lines: list[str] = []

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            if stripped.startswith('```'):
                if not in_code:
                    in_code = True
                    code_lang = stripped[3:].strip()
                    code_lines = []
                else:
                    in_code = False
                    p = self.doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25)
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(8.5)
                    if code_lang == 'mermaid':
                        note = self.doc.add_paragraph()
                        nr = note.add_run(
                            '[Diagrama Mermaid — consulte la versión Markdown o exporte desde GitHub]'
                        )
                        nr.italic = True
                        nr.font.size = Pt(9)
                        nr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                i += 1
                continue

            if in_code:
                code_lines.append(line)
                i += 1
                continue

            if not stripped:
                i += 1
                continue

            if stripped == '---':
                self.doc.add_paragraph('_' * 72)
                i += 1
                continue

            hm = re.match(r'^(#{1,4})\s+(.*)$', stripped)
            if hm:
                level = len(hm.group(1))
                self.doc.add_heading(hm.group(2).strip(), level=level)
                i += 1
                continue

            if is_table_row(stripped):
                rows, i = parse_table(lines, i)
                if rows:
                    self._add_table(rows)
                continue

            if re.match(r'^[-*]\s+', stripped):
                while i < len(lines) and re.match(r'^[-*]\s+', lines[i].strip()):
                    item = re.sub(r'^[-*]\s+', '', lines[i].strip())
                    p = self.doc.add_paragraph(style='List Bullet')
                    parse_inline(item, p)
                    i += 1
                continue

            if re.match(r'^\d+\.\s+', stripped):
                while i < len(lines) and re.match(r'^\d+\.\s+', lines[i].strip()):
                    item = re.sub(r'^\d+\.\s+', '', lines[i].strip())
                    p = self.doc.add_paragraph(style='List Number')
                    parse_inline(item, p)
                    i += 1
                continue

            if stripped.startswith('>'):
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.35)
                text = stripped.lstrip('> ').strip()
                run = p.add_run(text)
                run.italic = True
                i += 1
                continue

            p = self.doc.add_paragraph()
            parse_inline(stripped, p)
            i += 1

    def _add_table(self, rows: list[list[str]]) -> None:
        cols = max(len(r) for r in rows)
        table = self.doc.add_table(rows=len(rows), cols=cols)
        table.style = 'Table Grid'
        for ri, row in enumerate(rows):
            for ci in range(cols):
                cell = table.rows[ri].cells[ci]
                text = row[ci] if ci < len(row) else ''
                cell.text = ''
                p = cell.paragraphs[0]
                parse_inline(text, p)
                if ri == 0:
                    set_cell_shading(cell, 'E8EEF4')
                    for run in p.runs:
                        run.bold = True
        self.doc.add_paragraph()


def main() -> None:
    if len(sys.argv) < 2:
        print('Usage: python md-to-docx.py <input.md> [output.docx]')
        sys.exit(1)
    src = Path(sys.argv[1]).resolve()
    dst = Path(sys.argv[2]).resolve() if len(sys.argv) > 2 else src.with_suffix('.docx')

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    title = doc.add_heading(src.stem.replace('-', ' '), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    converter = MdToDocx(doc)
    converter.convert(src.read_text(encoding='utf-8'))

    doc.save(str(dst))
    print(f'OK: {dst}')


if __name__ == '__main__':
    main()
